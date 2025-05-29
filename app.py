# app.py
import streamlit as st
from openai import AzureOpenAI
import PyPDF2
import docx
import io
import json
import re
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import black, blue
import base64
from dotenv import load_dotenv
import os
from collections import Counter
import numpy as np

# Load environment variables
load_dotenv()

# Configure Streamlit page
st.set_page_config(
    page_title="Nagashree's ATS Resume Builder",
    page_icon="❤️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 3rem;
        font-weight: bold;
        text-align: center;
        color: #1f77b4;
        margin-bottom: 2rem;
    }
    .sub-header {
        font-size: 1.5rem;
        font-weight: bold;
        color: #333;
        margin: 1rem 0;
    }
    .metric-card {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #1f77b4;
    }
    .resume-preview {
        background-color: white;
        padding: 2rem;
        border-radius: 0.5rem;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        border: 1px solid #e0e0e0;
    }
    .success-box {
        background-color: #d4edda;
        border-color: #c3e6cb;
        color: #155724;
        padding: 1rem;
        border-radius: 0.5rem;
        border: 1px solid #c3e6cb;
    }
    .warning-box {
        background-color: #fff3cd;
        border-color: #ffeaa7;
        color: #856404;
        padding: 1rem;
        border-radius: 0.5rem;
        border: 1px solid #ffeaa7;
    }
    .ai-insight {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
            
    
    .footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        text-align: center;
        padding: 10px 0;
        font-size: 14px;
        font-weight: 500;
        box-shadow: 0 -2px 10px rgba(0,0,0,0.1);
        z-index: 999;
    }
    
    .footer-heart {
        color: #ff6b6b;
        font-size: 16px;
        animation: heartbeat 1.5s ease-in-out infinite;
    }
    
    @keyframes heartbeat {
        0% { transform: scale(1); }
        50% { transform: scale(1.1); }
        100% { transform: scale(1); }
    }
    
    .footer-name {
        color: #ffd93d;
        font-weight: bold;
        text-transform: uppercase;
        letter-spacing: 1px;
    }
    
    /* Add bottom padding to main content to avoid footer overlap */
    .main .block-container {
        padding-bottom: 60px;
    }
</style>
""", unsafe_allow_html=True)

class AIResumeOptimizer:
    def __init__(self):
        self.azure_client = None
        self.setup_azure_openai()
    
    
    
    def setup_azure_openai(self):
        """Initialize Azure OpenAI client"""
        try:
            endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
            api_key = os.getenv("AZURE_OPENAI_API_KEY")
            api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-01")
            
            if not endpoint or not api_key:
                st.error("Please set your Azure OpenAI credentials in the environment variables")
                st.stop()
            
            self.azure_client = AzureOpenAI(
                azure_endpoint=endpoint,
                api_key=api_key,
                api_version=api_version
            )
            
            self.deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")
            self.embeddings_deployment = os.getenv("AZURE_OPENAI_EMBEDDINGS_DEPLOYMENT_NAME")
            
            if not self.deployment_name:
                st.error("Please set AZURE_OPENAI_DEPLOYMENT_NAME in environment variables")
                st.stop()
                
        except Exception as e:
            st.error(f"Error setting up Azure OpenAI: {str(e)}")
            st.stop()
    
    def extract_text_from_pdf(self, pdf_file):
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            st.error(f"Error reading PDF: {str(e)}")
            return None
    
    def extract_text_from_docx(self, docx_file):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(docx_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            st.error(f"Error reading DOCX: {str(e)}")
            return None
    
    def ai_parse_resume(self, text):
        """Use AI to intelligently parse resume text"""
        try:
            prompt = f"""
You are an expert resume parser. Extract and structure the following resume text into a comprehensive JSON format.

Resume Text:
{text}

Please extract and organize the information into this exact JSON structure:
{{
    "personal_info": {{
        "name": "Full name",
        "email": "email@example.com",
        "phone": "phone number",
        "location": "city, state/country",
        "linkedin": "LinkedIn URL",
        "website": "personal website if any",
        "github": "GitHub profile if any"
    }},
    "summary": "Professional summary or objective statement",
    "experience": [
        {{
            "title": "Job title",
            "company": "Company name",
            "location": "Work location",
            "duration": "Start date - End date",
            "description": [
                "Achievement or responsibility 1",
                "Achievement or responsibility 2"
            ]
        }}
    ],
    "skills": {{
        "technical": ["skill1", "skill2"],
        "programming": ["language1", "language2"],
        "tools": ["tool1", "tool2"],
        "soft_skills": ["skill1", "skill2"]
    }},
    "education": [
        {{
            "degree": "Degree type and major",
            "school": "Institution name",
            "location": "School location",
            "year": "Graduation year",
            "gpa": "GPA if mentioned",
            "honors": "Any honors or distinctions"
        }}
    ],
    "projects": [
        {{
            "name": "Project name",
            "description": "Brief description",
            "technologies": ["tech1", "tech2"],
            "duration": "Project duration",
            "url": "Project URL if any"
        }}
    ],
    "certifications": [
        {{
            "name": "Certification name",
            "issuer": "Issuing organization",
            "date": "Date obtained",
            "expiry": "Expiry date if any"
        }}
    ],
    "languages": [
        {{
            "language": "Language name",
            "proficiency": "Proficiency level"
        }}
    ],
    "achievements": [
        "Achievement 1",
        "Achievement 2"
    ]
}}

Important: 
- Extract ALL information present in the resume
- If information is not available, use null or empty array
- Be thorough and accurate
- Maintain original formatting and content
- Return only valid JSON without any markdown formatting
"""

            response = self.azure_client.chat.completions.create(
                model=self.deployment_name,
                messages=[
                    {"role": "system", "content": "You are an expert resume parser. Return only valid JSON data."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=4000
            )
            
            content = response.choices[0].message.content.strip()
            
            # Clean up response
            if content.startswith('```json'):
                content = content[7:-3]
            elif content.startswith('```'):
                content = content[3:-3]
            
            return json.loads(content)
            
        except Exception as e:
            st.error(f"Error parsing resume with AI: {str(e)}")
            return self.fallback_parse_resume(text)
    
    def ai_extract_job_requirements(self, job_description):
        """Use AI to extract structured requirements from job description"""
        try:
            prompt = f"""
Analyze this job description and extract key requirements in structured format:

Job Description:
{job_description}

Extract and return as JSON:
{{
    "required_skills": ["skill1", "skill2"],
    "preferred_skills": ["skill1", "skill2"],
    "experience_level": "entry/mid/senior level",
    "education_requirements": ["requirement1", "requirement2"],
    "key_responsibilities": ["responsibility1", "responsibility2"],
    "industry": "industry name",
    "job_type": "full-time/part-time/contract",
    "keywords": ["keyword1", "keyword2"],
    "company_culture": ["value1", "value2"],
    "must_have_technologies": ["tech1", "tech2"],
    "nice_to_have_technologies": ["tech1", "tech2"]
}}

Return only valid JSON without markdown.
"""

            response = self.azure_client.chat.completions.create(
                model=self.deployment_name,
                messages=[
                    {"role": "system", "content": "You are an expert job requirement analyzer. Return only valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=2000
            )
            
            content = response.choices[0].message.content.strip()
            if content.startswith('```json'):
                content = content[7:-3]
            elif content.startswith('```'):
                content = content[3:-3]
            
            return json.loads(content)
            
        except Exception as e:
            st.error(f"Error analyzing job requirements: {str(e)}")
            return {}
    
    def ai_optimize_resume(self, resume_data, job_requirements, job_description):
        """Use AI to comprehensively optimize resume"""
        try:
            prompt = f"""
You are an expert ATS resume optimizer. Optimize this resume to perfectly match the job requirements while maintaining truthfulness.

Original Resume:
{json.dumps(resume_data, indent=2)}

Job Requirements:
{json.dumps(job_requirements, indent=2)}

Job Description:
{job_description}

Optimization Instructions:
1. Rewrite the professional summary to align with job requirements
2. Enhance experience descriptions with relevant keywords and quantified achievements
3. Reorganize and prioritize skills based on job requirements
4. Add missing relevant skills if they can be reasonably inferred from experience
5. Optimize formatting for ATS compatibility
6. Use action verbs and metrics where possible
7. Ensure keyword density matches job requirements
8. Maintain all factual information - DO NOT fabricate experience

Return the optimized resume in the same JSON structure with these improvements:
- Enhanced professional summary (3-4 lines)
- Optimized experience descriptions with relevant keywords
- Prioritized skills sections
- ATS-friendly formatting
- Strategic keyword placement

Return only the JSON structure without any markdown formatting.
"""

            response = self.azure_client.chat.completions.create(
                model=self.deployment_name,
                messages=[
                    {"role": "system", "content": "You are an expert ATS resume optimizer. Return only valid JSON data."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=4000
            )
            
            content = response.choices[0].message.content.strip()
            if content.startswith('```json'):
                content = content[7:-3]
            elif content.startswith('```'):
                content = content[3:-3]
            
            return json.loads(content)
            
        except Exception as e:
            st.error(f"Error optimizing resume: {str(e)}")
            return resume_data
    
    def ai_calculate_ats_score(self, resume_data, job_requirements):
        """Use AI to calculate comprehensive ATS score"""
        try:
            prompt = f"""
Analyze this resume against job requirements and calculate an ATS compatibility score.

Resume:
{json.dumps(resume_data, indent=2)}

Job Requirements:
{json.dumps(job_requirements, indent=2)}

Evaluate based on:
1. Keyword matching (30%)
2. Skills alignment (25%)
3. Experience relevance (20%)
4. Education requirements (10%)
5. Resume structure/formatting (10%)
6. Achievement quantification (5%)

Return analysis as JSON:
{{
    "overall_score": 85,
    "keyword_score": 80,
    "skills_score": 90,
    "experience_score": 85,
    "education_score": 75,
    "structure_score": 95,
    "achievement_score": 70,
    "matched_keywords": ["keyword1", "keyword2"],
    "missing_keywords": ["keyword1", "keyword2"],
    "strengths": ["strength1", "strength2"],
    "improvements": ["improvement1", "improvement2"],
    "ats_recommendations": ["rec1", "rec2"]
}}

Return only valid JSON.
"""

            response = self.azure_client.chat.completions.create(
                model=self.deployment_name,
                messages=[
                    {"role": "system", "content": "You are an ATS scoring expert. Return only valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=2000
            )
            
            content = response.choices[0].message.content.strip()
            if content.startswith('```json'):
                content = content[7:-3]
            elif content.startswith('```'):
                content = content[3:-3]
            
            return json.loads(content)
            
        except Exception as e:
            st.error(f"Error calculating ATS score: {str(e)}")
            return {"overall_score": 0, "error": str(e)}
    
    def ai_generate_cover_letter(self, resume_data, job_requirements, job_description, company_name=""):
        """Generate personalized cover letter"""
        try:
            prompt = f"""
Write a compelling, personalized cover letter based on this resume and job requirements.

Resume Data:
{json.dumps(resume_data, indent=2)}

Job Requirements:
{json.dumps(job_requirements, indent=2)}

Company: {company_name}
Job Description: {job_description}

Requirements:
- 3-4 paragraphs
- Professional tone
- Highlight relevant experience
- Show enthusiasm for the role
- Include specific achievements
- Match company culture if mentioned
- ATS-friendly format

Return as plain text, ready to use.
"""

            response = self.azure_client.chat.completions.create(
                model=self.deployment_name,
                messages=[
                    {"role": "system", "content": "You are an expert cover letter writer."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.4,
                max_tokens=1500
            )
            
            return response.choices[0].message.content.strip()
            
        except Exception as e:
            st.error(f"Error generating cover letter: {str(e)}")
            return ""
    
    def ai_suggest_improvements(self, resume_data, ats_analysis):
        """Get AI-powered improvement suggestions"""
        try:
            prompt = f"""
Based on this resume and ATS analysis, provide specific, actionable improvement suggestions.

Resume:
{json.dumps(resume_data, indent=2)}

ATS Analysis:
{json.dumps(ats_analysis, indent=2)}

Provide detailed suggestions in categories:
1. Content improvements
2. Keyword optimization
3. Formatting enhancements
4. Skills development
5. Experience enhancement

Return as JSON:
{{
    "content_improvements": ["suggestion1", "suggestion2"],
    "keyword_optimization": ["suggestion1", "suggestion2"],
    "formatting_enhancements": ["suggestion1", "suggestion2"],
    "skills_development": ["suggestion1", "suggestion2"],
    "experience_enhancement": ["suggestion1", "suggestion2"],
    "quick_wins": ["quick fix 1", "quick fix 2"],
    "long_term_goals": ["goal1", "goal2"]
}}
"""

            response = self.azure_client.chat.completions.create(
                model=self.deployment_name,
                messages=[
                    {"role": "system", "content": "You are a career development expert."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=2000
            )
            
            content = response.choices[0].message.content.strip()
            if content.startswith('```json'):
                content = content[7:-3]
            elif content.startswith('```'):
                content = content[3:-3]
            
            return json.loads(content)
            
        except Exception as e:
            st.error(f"Error generating suggestions: {str(e)}")
            return {}
    
    def fallback_parse_resume(self, text):
        """Fallback parser if AI fails"""
        return {
            "personal_info": {"name": "", "email": "", "phone": ""},
            "summary": "",
            "experience": [],
            "skills": {"technical": [], "programming": [], "tools": [], "soft_skills": []},
            "education": [],
            "projects": [],
            "certifications": [],
            "languages": [],
            "achievements": []
        }
    
    def generate_pdf_resume(self, resume_data):
        """Generate professional PDF resume with proper alignment and formatting"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=letter, 
            rightMargin=40, 
            leftMargin=40,
            topMargin=40, 
            bottomMargin=40
        )
        
        # Get styles and create custom ones
        styles = getSampleStyleSheet()
        story = []
        
        # Enhanced custom styles with better alignment
        name_style = ParagraphStyle(
            'NameStyle',
            parent=styles['Title'],
            fontSize=20,
            fontName='Helvetica-Bold',
            textColor=black,
            alignment=1,  # Center align for name
            spaceAfter=6,
            spaceBefore=0
        )
        
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Normal'],
            fontSize=12,
            fontName='Helvetica',
            textColor=black,
            alignment=1,  # Center align for title
            spaceAfter=12,
            spaceBefore=0
        )
        
        contact_style = ParagraphStyle(
            'ContactStyle',
            parent=styles['Normal'],
            fontSize=10,
            fontName='Helvetica',
            textColor=black,
            alignment=0,  # Left align
            spaceAfter=2,
            spaceBefore=0,
            leftIndent=0,
            rightIndent=0
        )
        
        section_heading_style = ParagraphStyle(
            'SectionHeading',
            parent=styles['Heading2'],
            fontSize=13,
            fontName='Helvetica-Bold',
            textColor=black,
            alignment=0,  # Left align
            spaceAfter=10,
            spaceBefore=16,
            borderWidth=1,
            borderColor=black,
            borderPadding=2,
            keepWithNext=True,  # Prevent orphaned headers
            pageBreakBefore=0,  # Allow page breaks before if needed
            splitLongWords=False
        )
        
        subsection_style = ParagraphStyle(
            'SubsectionStyle',
            parent=styles['Normal'],
            fontSize=11,
            fontName='Helvetica-Bold',
            textColor=black,
            alignment=0,
            spaceAfter=4,
            spaceBefore=8,
            keepWithNext=True
        )
        
        body_style = ParagraphStyle(
            'BodyStyle',
            parent=styles['Normal'],
            fontSize=10,
            fontName='Helvetica',
            textColor=black,
            alignment=4,  # Justify
            spaceAfter=4,
            spaceBefore=0,
            leftIndent=0,
            rightIndent=0
        )
        
        bullet_style = ParagraphStyle(
            'BulletStyle',
            parent=styles['Normal'],
            fontSize=10,
            fontName='Helvetica',
            textColor=black,
            alignment=0,
            spaceAfter=3,
            spaceBefore=0,
            leftIndent=15,
            bulletIndent=10,
            bulletFontName='Symbol'
        )
        
        # Helper function to add section with proper formatting and page break protection
        def add_section_header(text):
            from reportlab.platypus import KeepTogether, PageBreak
            
            # Create header paragraph with page break protection
            header_para = Paragraph(f'<b>{text.upper()}</b>', section_heading_style)
            
            # Add a thin line under section headers
            from reportlab.platypus import HRFlowable
            hr_line = HRFlowable(width="100%", thickness=1, lineCap='round', color=black)
            spacer = Spacer(1, 4)
            
            # Group header, line, and spacer together to prevent separation
            header_group = KeepTogether([header_para, hr_line, spacer])
            story.append(header_group)
        
        # HEADER SECTION WITH IMPROVED ALIGNMENT
        personal_info = resume_data.get('personal_info', {})
        
        # Name - centered and prominent
        if personal_info.get('name'):
            story.append(Paragraph(f"<b>{personal_info['name'].upper()}</b>", name_style))
        
        # Professional title - centered
        title_text = "AI Engineer | Full Stack Developer | DevOps Specialist"
        if resume_data.get('summary'):
            summary = resume_data['summary']
            # Extract professional roles from summary if available
            if any(role in summary.lower() for role in ['engineer', 'developer', 'analyst', 'specialist']):
                title_text = self.extract_professional_title(summary)
        
        story.append(Paragraph(title_text, title_style))
        story.append(Spacer(1, 8))
        
        # Contact information in properly aligned table
        contact_data = self.create_contact_table(personal_info)
        if contact_data:
            from reportlab.platypus import Table, TableStyle
            from reportlab.lib import colors
            
            # Create balanced contact table
            contact_table = Table(contact_data, colWidths=[2.8*inch, 2.8*inch])
            contact_table.setStyle(TableStyle([
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),   # Left column left-aligned
                ('ALIGN', (1, 0), (1, -1), 'RIGHT'),  # Right column right-aligned
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.white),  # Invisible grid for spacing
            ]))
            story.append(contact_table)
        
        story.append(Spacer(1, 16))
        
        # PROFESSIONAL SUMMARY with proper formatting
        if resume_data.get('summary'):
            add_section_header("PROFESSIONAL SUMMARY")
            summary_para = Paragraph(resume_data['summary'], body_style)
            story.append(summary_para)
            story.append(Spacer(1, 12))
        
        # CORE COMPETENCIES with left-aligned layout, page break protection, and proper text wrapping
        skills = resume_data.get('skills', {})
        if skills:
            add_section_header("CORE COMPETENCIES")
            
            # Organize skills in a clean table format with Paragraph objects
            skill_table_data = self.organize_skills_table_with_paragraphs(skills, body_style)
            if skill_table_data:
                # Fixed column widths for better alignment - adjust for wrapping
                skills_table = Table(skill_table_data, colWidths=[1.5*inch, 4.2*inch])
                skills_table.setStyle(TableStyle([
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),  # Category names bold
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),              # Top align for multi-line content
                    ('ALIGN', (0, 0), (0, -1), 'LEFT'),               # Left-align category labels
                    ('ALIGN', (1, 0), (1, -1), 'LEFT'),               # Left-align skills
                    ('LEFTPADDING', (0, 0), (0, -1), 0),              # No left padding for categories
                    ('RIGHTPADDING', (0, 0), (0, -1), 10),            # Right padding for categories
                    ('LEFTPADDING', (1, 0), (1, -1), 10),             # Left padding for skills
                    ('RIGHTPADDING', (1, 0), (1, -1), 5),             # Small right padding for skills
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ('GRID', (0, 0), (-1, -1), 0, colors.white),      # No visible grid
                ]))
                # Ensure table stays with header
                skills_table.keepWithNext = True
                story.append(skills_table)
            
            story.append(Spacer(1, 12))
        
        # WORK EXPERIENCE with consistent formatting and page break protection
        if resume_data.get('experience'):
            add_section_header("PROFESSIONAL EXPERIENCE")
            
            for i, exp in enumerate(resume_data['experience']):
                # Create experience block that stays together
                exp_elements = []
                
                # Job title in bold
                if exp.get('title'):
                    title_para = Paragraph(f"<b>{exp['title']}</b>", subsection_style)
                    exp_elements.append(title_para)
                
                # Company, location, and duration in organized format
                exp_info = self.format_experience_header(exp)
                if exp_info:
                    company_para = Paragraph(exp_info, body_style)
                    exp_elements.append(company_para)
                
                # Job responsibilities with proper bullet formatting
                if exp.get('description'):
                    if isinstance(exp['description'], list):
                        for desc in exp['description']:
                            if desc.strip():
                                bullet_para = Paragraph(f"• {desc}", bullet_style)
                                exp_elements.append(bullet_para)
                    else:
                        bullet_para = Paragraph(f"• {exp['description']}", bullet_style)
                        exp_elements.append(bullet_para)
                
                # Keep experience block together
                from reportlab.platypus import KeepTogether
                if exp_elements:
                    exp_block = KeepTogether(exp_elements)
                    story.append(exp_block)
                
                # Add spacing between experiences
                if i < len(resume_data['experience']) - 1:
                    story.append(Spacer(1, 10))
            
            story.append(Spacer(1, 12))
        
        # KEY PROJECTS with enhanced formatting and page break protection
        if resume_data.get('projects'):
            add_section_header("KEY PROJECTS")
            
            for i, project in enumerate(resume_data['projects']):
                # Create project block that stays together
                project_elements = []
                
                # Project name and duration
                project_header = self.format_project_header(project)
                if project_header:
                    header_para = Paragraph(project_header, subsection_style)
                    project_elements.append(header_para)
                
                # Project description and details
                if project.get('description'):
                    if isinstance(project['description'], list):
                        for desc in project['description']:
                            if desc.strip():
                                desc_para = Paragraph(f"• {desc}", bullet_style)
                                project_elements.append(desc_para)
                    else:
                        desc_para = Paragraph(f"• {project['description']}", bullet_style)
                        project_elements.append(desc_para)
                
                # Technologies with proper formatting
                if project.get('technologies'):
                    tech_text = self.format_technologies(project['technologies'])
                    tech_para = Paragraph(f"• <b>Technologies:</b> {tech_text}", bullet_style)
                    project_elements.append(tech_para)
                
                # Keep project block together
                from reportlab.platypus import KeepTogether
                if project_elements:
                    project_block = KeepTogether(project_elements)
                    story.append(project_block)
                
                if i < len(resume_data['projects']) - 1:
                    story.append(Spacer(1, 8))
            
            story.append(Spacer(1, 12))
        
        # EDUCATION with consistent alignment and page break protection
        if resume_data.get('education'):
            add_section_header("EDUCATION")
            
            for edu in resume_data['education']:
                # Create education block that stays together
                edu_elements = []
                
                # Degree in bold
                if edu.get('degree'):
                    degree_para = Paragraph(f"<b>{edu['degree']}</b>", subsection_style)
                    edu_elements.append(degree_para)
                
                # School and details
                edu_info = self.format_education_info(edu)
                if edu_info:
                    school_para = Paragraph(edu_info, body_style)
                    edu_elements.append(school_para)
                
                # Additional details (GPA, honors, etc.)
                edu_details = self.format_education_details(edu)
                if edu_details:
                    for detail in edu_details:
                        detail_para = Paragraph(detail, body_style)
                        edu_elements.append(detail_para)
                
                # Keep education block together
                from reportlab.platypus import KeepTogether
                if edu_elements:
                    edu_block = KeepTogether(edu_elements)
                    story.append(edu_block)
                
                story.append(Spacer(1, 6))
        
        # CERTIFICATIONS with organized layout
        if resume_data.get('certifications'):
            add_section_header("CERTIFICATIONS")
            
            for cert in resume_data['certifications']:
                cert_text = self.format_certification(cert)
                if cert_text:
                    story.append(Paragraph(f"• {cert_text}", bullet_style))
            
            story.append(Spacer(1, 8))
        
        # ACHIEVEMENTS with proper formatting
        if resume_data.get('achievements'):
            add_section_header("ACHIEVEMENTS AND RECOGNITION")
            
            for achievement in resume_data['achievements']:
                story.append(Paragraph(f"• {achievement}", bullet_style))
        
        # Build the PDF
        doc.build(story)
        buffer.seek(0)
        return buffer

    def create_contact_table(self, personal_info):
        """Create properly formatted contact information table"""
        contact_left = []
        contact_right = []
        
        # Organize contact info in two balanced columns
        if personal_info.get('location'):
            contact_left.append(f"Location: {personal_info['location']}")
        if personal_info.get('linkedin'):
            linkedin_display = personal_info['linkedin'].replace('https://', '').replace('http://', '')
            contact_left.append(f"LinkedIn: {linkedin_display}")
        
        if personal_info.get('phone'):
            contact_right.append(f"Mobile: {personal_info['phone']}")
        if personal_info.get('email'):
            contact_right.append(f"Email: {personal_info['email']}")
        if personal_info.get('github'):
            github_display = personal_info['github'].replace('https://', '').replace('http://', '')
            contact_right.append(f"GitHub: {github_display}")
        
        # Create balanced table data
        contact_data = []
        max_items = max(len(contact_left), len(contact_right))
        
        for i in range(max_items):
            left_item = contact_left[i] if i < len(contact_left) else ""
            right_item = contact_right[i] if i < len(contact_right) else ""
            contact_data.append([left_item, right_item])
        
        return contact_data if contact_data else None

    def organize_skills_table_with_paragraphs(self, skills, body_style):
        """Organize skills into table format using Paragraph objects for proper wrapping"""
        skill_table_data = []
        
        # Create a style for skills with wrapping
        skills_style = ParagraphStyle(
            'SkillsStyle',
            parent=body_style,
            fontSize=10,
            fontName='Helvetica',
            textColor=black,
            alignment=0,  # Left align
            spaceAfter=0,
            spaceBefore=0,
            leftIndent=0,
            rightIndent=0,
            wordWrap='LTR'
        )
        
        # Create a style for category labels
        category_style = ParagraphStyle(
            'CategoryStyle',
            parent=body_style,
            fontSize=10,
            fontName='Helvetica-Bold',
            textColor=black,
            alignment=0,  # Left align
            spaceAfter=0,
            spaceBefore=0
        )
        
        if isinstance(skills, dict):
            # Define skill categories and their display order
            skill_categories = {
                'programming': 'Programming',
                'technical': 'Technical',
                'ai_ml': 'AI/ML',
                'cloud': 'Cloud',
                'tools': 'Tools & Platforms',
                'databases': 'Databases',
                'frameworks': 'Frameworks',
                'soft_skills': 'Soft Skills'
            }
            
            # Categorize skills intelligently
            categorized_skills = self.categorize_skills_intelligently(skills)
            
            for category_key, category_name in skill_categories.items():
                if categorized_skills.get(category_key):
                    # Create category paragraph
                    category_para = Paragraph(f"{category_name}:", category_style)
                    
                    # Format skills with proper wrapping using Paragraph
                    skills_text = ", ".join(categorized_skills[category_key])
                    skills_para = Paragraph(skills_text, skills_style)
                    
                    skill_table_data.append([category_para, skills_para])
        
        elif isinstance(skills, list):
            category_para = Paragraph("Technical Skills:", category_style)
            skills_text = ", ".join(skills)
            skills_para = Paragraph(skills_text, skills_style)
            skill_table_data.append([category_para, skills_para])
        
        return skill_table_data

    def organize_skills_table(self, skills):
        """Organize skills into a clean table format with proper alignment and text wrapping"""
        skill_table_data = []
        
        if isinstance(skills, dict):
            # Define skill categories and their display order
            skill_categories = {
                'programming': 'Programming',
                'technical': 'Technical',
                'ai_ml': 'AI/ML',
                'cloud': 'Cloud',
                'tools': 'Tools & Platforms',
                'databases': 'Databases',
                'frameworks': 'Frameworks',
                'soft_skills': 'Soft Skills'
            }
            
            # Categorize skills intelligently
            categorized_skills = self.categorize_skills_intelligently(skills)
            
            for category_key, category_name in skill_categories.items():
                if categorized_skills.get(category_key):
                    # Format skills with proper wrapping
                    skills_text = self.format_skills_with_wrapping(categorized_skills[category_key])
                    # Ensure consistent formatting with colon
                    skill_table_data.append([f"{category_name}:", skills_text])
        
        elif isinstance(skills, list):
            skills_text = self.format_skills_with_wrapping(skills)
            skill_table_data.append(["Technical Skills:", skills_text])
        
        return skill_table_data

    def format_skills_with_wrapping(self, skills_list, max_line_length=80):
        """Format skills list with proper line breaks to prevent overflow"""
        if not skills_list:
            return ""
        
        # Join skills with commas
        skills_text = ", ".join(skills_list)
        
        # If the text is short enough, return as is
        if len(skills_text) <= max_line_length:
            return skills_text
        
        # Break into multiple lines while keeping skills together
        lines = []
        current_line = ""
        
        for skill in skills_list:
            # Check if adding this skill would exceed line length
            if current_line and len(current_line + ", " + skill) > max_line_length:
                # Add current line and start new one
                lines.append(current_line)
                current_line = skill
            else:
                # Add to current line
                if current_line:
                    current_line += ", " + skill
                else:
                    current_line = skill
        
        # Add the last line
        if current_line:
            lines.append(current_line)
        
        # Join lines with line breaks
        return "<br/>".join(lines)

    def categorize_skills_intelligently(self, skills):
        """Intelligently categorize skills based on keywords"""
        categorized = {
            'programming': [],
            'ai_ml': [],
            'cloud': [],
            'tools': [],
            'databases': [],
            'frameworks': [],
            'technical': []
        }
        
        # Define keyword mappings
        category_keywords = {
            'programming': ['python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'go', 'rust', 'php'],
            'ai_ml': ['ai', 'ml', 'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'opencv', 'nlp', 'computer vision'],
            'cloud': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'cloud'],
            'databases': ['sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'database'],
            'frameworks': ['react', 'angular', 'vue', 'node', 'express', 'django', 'flask', 'spring'],
            'tools': ['git', 'jenkins', 'ci/cd', 'linux', 'bash']
        }
        
        # Flatten all skills
        all_skills = []
        for skill_list in skills.values():
            if isinstance(skill_list, list):
                all_skills.extend(skill_list)
        
        # Categorize each skill
        for skill in all_skills:
            skill_lower = skill.lower()
            categorized_skill = False
            
            for category, keywords in category_keywords.items():
                if any(keyword in skill_lower for keyword in keywords):
                    categorized[category].append(skill)
                    categorized_skill = True
                    break
            
            if not categorized_skill:
                categorized['technical'].append(skill)
        
        # Remove empty categories and duplicates
        return {k: list(set(v)) for k, v in categorized.items() if v}

    def format_experience_header(self, exp):
        """Format experience header with company, location, and duration"""
        parts = []
        
        if exp.get('company'):
            parts.append(f"<b>{exp['company']}</b>")
        
        # Add location and duration on same line, properly spaced
        details = []
        if exp.get('location'):
            details.append(exp['location'])
        if exp.get('duration'):
            details.append(exp['duration'])
        
        if details:
            parts.append(" | ".join(details))
        
        return " | ".join(parts) if parts else None

    def format_project_header(self, project):
        """Format project header with name and duration"""
        if not project.get('name'):
            return None
        
        header = f"<b>{project['name']}</b>"
        if project.get('duration'):
            header += f" | <i>{project['duration']}</i>"
        
        return header

    def format_technologies(self, technologies):
        """Format technologies list"""
        if isinstance(technologies, list):
            return ", ".join(technologies)
        return str(technologies)

    def format_education_info(self, edu):
        """Format education institution and year"""
        parts = []
        
        if edu.get('school'):
            parts.append(edu['school'])
        if edu.get('location'):
            parts.append(edu['location'])
        if edu.get('year'):
            parts.append(str(edu['year']))
        
        return " | ".join(parts) if parts else None

    def format_education_details(self, edu):
        """Format additional education details"""
        details = []
        
        if edu.get('gpa'):
            details.append(f"CGPA: {edu['gpa']}")
        if edu.get('percentage'):
            details.append(f"Percentage: {edu['percentage']}%")
        if edu.get('honors'):
            details.append(f"Honors: {edu['honors']}")
        
        return details

    def format_certification(self, cert):
        """Format certification information"""
        if isinstance(cert, dict):
            parts = []
            if cert.get('name'):
                parts.append(cert['name'])
            if cert.get('issuer'):
                parts.append(f"by {cert['issuer']}")
            if cert.get('date'):
                parts.append(f"({cert['date']})")
            return " ".join(parts)
        
        return str(cert)

    def extract_professional_title(self, summary):
        """Extract professional title from summary"""
        # Default title
        default_title = "AI Engineer | Full Stack Developer | DevOps Specialist"
        
        # Try to extract roles from summary
        role_keywords = {
            'ai': 'AI Engineer',
            'machine learning': 'ML Engineer', 
            'data scientist': 'Data Scientist',
            'full stack': 'Full Stack Developer',
            'backend': 'Backend Developer',
            'frontend': 'Frontend Developer',
            'devops': 'DevOps Engineer',
            'cloud': 'Cloud Engineer'
        }
        
        summary_lower = summary.lower()
        found_roles = []
        
        for keyword, role in role_keywords.items():
            if keyword in summary_lower:
                found_roles.append(role)
        
        if found_roles:
            return " | ".join(found_roles[:3])  # Limit to 3 roles
        
        return default_title

    def generate_docx_resume(self, resume_data):
        """Generate Word document resume"""
        try:
            doc = docx.Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = docx.shared.Inches(0.5)
                section.bottom_margin = docx.shared.Inches(0.5)
                section.left_margin = docx.shared.Inches(0.75)
                section.right_margin = docx.shared.Inches(0.75)
            
            personal_info = resume_data.get('personal_info', {})
            
            # Header with name
            if personal_info.get('name'):
                name_para = doc.add_paragraph()
                name_run = name_para.add_run(personal_info['name'].upper())
                name_run.font.size = docx.shared.Pt(18)
                name_run.font.bold = True
                name_para.alignment = docx.enum.text.WS_ALIGN_LEFT
            
            # Professional title
            title_para = doc.add_paragraph("AI Engineer | Full Stack Developer | DevOps Specialist")
            title_para.alignment = docx.enum.text.WS_ALIGN_LEFT
            
            # Contact information
            if any(personal_info.values()):
                # Create a table for contact info
                contact_table = doc.add_table(rows=3, cols=2)
                contact_table.style = 'Table Grid'
                
                # Left column
                if personal_info.get('linkedin'):
                    contact_table.cell(0, 0).text = f"LinkedIn: {personal_info['linkedin']}"
                if personal_info.get('location'):
                    contact_table.cell(1, 0).text = f"Location: {personal_info['location']}"
                
                # Right column
                if personal_info.get('github'):
                    contact_table.cell(0, 1).text = f"GitHub: {personal_info['github']}"
                if personal_info.get('phone'):
                    contact_table.cell(1, 1).text = f"Mob: {personal_info['phone']}"
                if personal_info.get('email'):
                    contact_table.cell(2, 1).text = f"E-mail: {personal_info['email']}"
            
            # Professional Summary
            if resume_data.get('summary'):
                doc.add_paragraph()
                summary_heading = doc.add_paragraph()
                summary_run = summary_heading.add_run("PROFESSIONAL SUMMARY")
                summary_run.font.bold = True
                summary_run.underline = True
                doc.add_paragraph(resume_data['summary'])
            
            # Core Competencies
            skills = resume_data.get('skills', {})
            if skills:
                doc.add_paragraph()
                skills_heading = doc.add_paragraph()
                skills_run = skills_heading.add_run("CORE COMPETENCIES")
                skills_run.font.bold = True
                skills_run.underline = True
                
                tech_para = doc.add_paragraph()
                tech_run = tech_para.add_run("Technical Skills:")
                tech_run.font.bold = True
                
                if isinstance(skills, dict):
                    for category, skill_list in skills.items():
                        if skill_list:
                            category_para = doc.add_paragraph()
                            category_run = category_para.add_run(f"- {category.replace('_', ' ').title()}: ")
                            category_run.font.bold = True
                            category_para.add_run(", ".join(skill_list))
                elif isinstance(skills, list):
                    skills_para = doc.add_paragraph()
                    skills_para.add_run("- Technical Skills: ")
                    skills_para.add_run(", ".join(skills))
            
            # Work Experience
            if resume_data.get('experience'):
                doc.add_paragraph()
                exp_heading = doc.add_paragraph()
                exp_run = exp_heading.add_run("WORK EXPERIENCE")
                exp_run.font.bold = True
                exp_run.underline = True
                
                for exp in resume_data['experience']:
                    if exp.get('title'):
                        title_para = doc.add_paragraph()
                        title_run = title_para.add_run(exp['title'])
                        title_run.font.bold = True
                    
                    if exp.get('company'):
                        company_para = doc.add_paragraph()
                        company_text = exp['company']
                        if exp.get('duration'):
                            company_text += f" | {exp['duration']}"
                        company_run = company_para.add_run(company_text)
                        company_run.font.bold = True
                    
                    if exp.get('description'):
                        if isinstance(exp['description'], list):
                            for desc in exp['description']:
                                if desc.strip():
                                    desc_para = doc.add_paragraph(f"- {desc}")
                        else:
                            desc_para = doc.add_paragraph(f"- {exp['description']}")
            
            # Projects
            if resume_data.get('projects'):
                doc.add_paragraph()
                proj_heading = doc.add_paragraph()
                proj_run = proj_heading.add_run("KEY PROJECTS")
                proj_run.font.bold = True
                proj_run.underline = True
                
                for project in resume_data['projects']:
                    if project.get('name'):
                        proj_para = doc.add_paragraph()
                        proj_name = project['name']
                        if project.get('duration'):
                            proj_name += f" | {project['duration']}"
                        proj_run = proj_para.add_run(proj_name)
                        proj_run.font.bold = True
                    
                    if project.get('description'):
                        if isinstance(project['description'], list):
                            for desc in project['description']:
                                if desc.strip():
                                    doc.add_paragraph(f"- {desc}")
                        else:
                            doc.add_paragraph(f"- {project['description']}")
                    
                    if project.get('technologies'):
                        tech_para = doc.add_paragraph()
                        tech_para.add_run("- Technologies: ").font.bold = True
                        if isinstance(project['technologies'], list):
                            tech_para.add_run(", ".join(project['technologies']))
                        else:
                            tech_para.add_run(project['technologies'])
            
            # Education
            if resume_data.get('education'):
                doc.add_paragraph()
                edu_heading = doc.add_paragraph()
                edu_run = edu_heading.add_run("EDUCATION")
                edu_run.font.bold = True
                edu_run.underline = True
                
                for edu in resume_data['education']:
                    if edu.get('degree'):
                        degree_para = doc.add_paragraph()
                        degree_run = degree_para.add_run(edu['degree'])
                        degree_run.font.bold = True
                    
                    if edu.get('school'):
                        school_text = edu['school']
                        if edu.get('year'):
                            school_text += f" | {edu['year']}"
                        doc.add_paragraph(school_text)
                    
                    if edu.get('gpa'):
                        doc.add_paragraph(f"CGPA: {edu['gpa']}")
                    elif edu.get('percentage'):
                        doc.add_paragraph(f"Percentage: {edu['percentage']}")
            
            # Certifications
            if resume_data.get('certifications'):
                doc.add_paragraph()
                cert_heading = doc.add_paragraph()
                cert_run = cert_heading.add_run("CERTIFICATIONS")
                cert_run.font.bold = True
                cert_run.underline = True
                
                for cert in resume_data['certifications']:
                    if isinstance(cert, dict):
                        cert_text = cert.get('name', '')
                        if cert.get('issuer'):
                            cert_text += f" - {cert['issuer']}"
                        if cert.get('date'):
                            cert_text += f" | {cert['date']}"
                        doc.add_paragraph(f"- {cert_text}")
                    else:
                        doc.add_paragraph(f"- {cert}")
            
            # Achievements
            if resume_data.get('achievements'):
                doc.add_paragraph()
                achv_heading = doc.add_paragraph()
                achv_run = achv_heading.add_run("ACHIEVEMENTS AND RECOGNITION")
                achv_run.font.bold = True
                achv_run.underline = True
                
                for achievement in resume_data['achievements']:
                    doc.add_paragraph(f"- {achievement}")
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            st.error(f"Error generating Word document: {str(e)}")
            return None

def main():
    st.markdown('<h1 class="main-header">❤️ Resume Builder for Nagshree ❤️</h1>', unsafe_allow_html=True)
    st.markdown('<p style="text-align: center; font-size: 1.2rem; color: #666;">Advanced Azure OpenAI Resume Optimization</p>', unsafe_allow_html=True)
    
    # Initialize optimizer
    optimizer = AIResumeOptimizer()

    
    
    # Sidebar
    with st.sidebar:
        st.markdown("### 🚀 AI Features")
        st.markdown("""
        - 🧠 **AI Resume Parsing** - Intelligent extraction
        - 📊 **AI ATS Scoring** - Comprehensive analysis  
        - ✨ **AI Optimization** - Smart enhancements
        - 📝 **AI Cover Letter** - Personalized generation
        - 💡 **AI Suggestions** - Expert recommendations
        """)
        
        st.markdown("### ⚙️ Azure OpenAI Status")
        if optimizer.azure_client:
            st.success("✅ Azure OpenAI Connected")
            st.info(f"Model: {optimizer.deployment_name}")
        else:
            st.error("❌ Azure OpenAI Not Connected")
        
        st.markdown("### 📋 Process")
        st.markdown("""
        1. Upload resume (PDF/DOCX)
        2. Paste job description
        3. AI analyzes & optimizes
        4. Review AI insights
        5. Download optimized resume
        """)
    
    # Main content
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown('<h2 class="sub-header">📄 Upload Resume</h2>', unsafe_allow_html=True)
        
        uploaded_file = st.file_uploader(
            "Choose your resume file",
            type=['pdf', 'docx'],
            help="Upload your current resume - AI will parse it intelligently"
        )
        
        if uploaded_file is not None:
            # Extract text
            if uploaded_file.type == "application/pdf":
                extracted_text = optimizer.extract_text_from_pdf(uploaded_file)
            else:
                extracted_text = optimizer.extract_text_from_docx(uploaded_file)
            
            if extracted_text:
                with st.spinner("🧠 AI is parsing your resume..."):
                    resume_data = optimizer.ai_parse_resume(extracted_text)
                
                st.success("✅ Resume parsed by AI!")
                st.session_state['resume_data'] = resume_data
                
                # AI-powered summary
                with st.expander("🤖 AI Analysis Summary"):
                    col_a, col_b, col_c = st.columns(3)
                    with col_a:
                        exp_count = len(resume_data.get('experience', []))
                        st.metric("Experience Items", exp_count)
                    with col_b:
                        all_skills = []
                        skills = resume_data.get('skills', {})
                        for skill_list in skills.values():
                            all_skills.extend(skill_list)
                        st.metric("Skills Found", len(all_skills))
                    with col_c:
                        edu_count = len(resume_data.get('education', []))
                        st.metric("Education Items", edu_count)
        
        # Job Description
        st.markdown('<h2 class="sub-header">🎯 Job Description</h2>', unsafe_allow_html=True)
        
        job_description = st.text_area(
            "Paste the complete job description",
            height=300,
            placeholder="Paste the job description here - AI will analyze requirements..."
        )
        
        company_name = st.text_input("Company Name (optional)", placeholder="e.g., Google, Microsoft")
        
        if job_description:
            st.session_state['job_description'] = job_description
            st.session_state['company_name'] = company_name
        
        # AI Optimization Button
        if st.button("🚀 AI Optimize Resume", type="primary", use_container_width=True):
            if 'resume_data' not in st.session_state:
                st.error("❌ Please upload a resume first")
            elif not job_description:
                st.error("❌ Please enter a job description")
            else:
                with st.spinner("🤖 AI is analyzing and optimizing..."):
                    # AI analyze job requirements
                    job_requirements = optimizer.ai_extract_job_requirements(job_description)
                    st.session_state['job_requirements'] = job_requirements
                    
                    # AI optimize resume
                    optimized_resume = optimizer.ai_optimize_resume(
                        st.session_state['resume_data'], 
                        job_requirements,
                        job_description
                    )
                    st.session_state['optimized_resume'] = optimized_resume
                    
                    # AI calculate ATS score
                    ats_analysis = optimizer.ai_calculate_ats_score(
                        optimized_resume,
                        job_requirements
                    )
                    st.session_state['ats_analysis'] = ats_analysis
                    
                    # AI generate suggestions
                    suggestions = optimizer.ai_suggest_improvements(
                        optimized_resume,
                        ats_analysis
                    )
                    st.session_state['ai_suggestions'] = suggestions
                
                st.success(f"✅ AI Optimization Complete! Score: {ats_analysis.get('overall_score', 0)}%")
    
    with col2:
        st.markdown('<h2 class="sub-header">👁️ Resume Preview</h2>', unsafe_allow_html=True)
        
        if 'optimized_resume' in st.session_state:
            tab1, tab2 = st.tabs(["🎯 AI Optimized", "📄 Original"])
            
            with tab1:
                display_resume_preview(st.session_state['optimized_resume'])
            with tab2:
                display_resume_preview(st.session_state['resume_data'])
        elif 'resume_data' in st.session_state:
            display_resume_preview(st.session_state['resume_data'])
        else:
            st.info("📤 Upload a resume to see AI-powered preview")
    
    # AI Results Section
    if 'ats_analysis' in st.session_state:
        st.markdown("---")
        st.markdown('<h2 class="sub-header">🤖 AI Analysis Results</h2>', unsafe_allow_html=True)
        
        ats_analysis = st.session_state['ats_analysis']
        
        # AI Score Dashboard
        col1, col2, col3, col4, col5 = st.columns(5)
        with col1:
            score = ats_analysis.get('overall_score', 0)
            color = "🟢" if score >= 90 else "🟡" if score >= 75 else "🔴"
            st.metric("AI ATS Score", f"{score}%", delta=color)
        with col2:
            st.metric("Keyword Match", f"{ats_analysis.get('keyword_score', 0)}%")
        with col3:
            st.metric("Skills Alignment", f"{ats_analysis.get('skills_score', 0)}%")
        with col4:
            st.metric("Experience Match", f"{ats_analysis.get('experience_score', 0)}%")
        with col5:
            st.metric("Structure Score", f"{ats_analysis.get('structure_score', 0)}%")
        
        # AI Insights
        col_insight1, col_insight2 = st.columns(2)
        
        with col_insight1:
            st.markdown('<div class="ai-insight">', unsafe_allow_html=True)
            st.markdown("### 🎯 **AI Strengths Analysis**")
            for strength in ats_analysis.get('strengths', [])[:3]:
                st.markdown(f"✅ {strength}")
            st.markdown('</div>', unsafe_allow_html=True)
        
        with col_insight2:
            st.markdown('<div class="ai-insight">', unsafe_allow_html=True)
            st.markdown("### 🔧 **AI Improvement Areas**")
            for improvement in ats_analysis.get('improvements', [])[:3]:
                st.markdown(f"🔧 {improvement}")
            st.markdown('</div>', unsafe_allow_html=True)
        
        # Detailed AI Analysis
        with st.expander("🔍 Detailed AI Analysis"):
            analysis_col1, analysis_col2 = st.columns(2)
            
            with analysis_col1:
                st.markdown("**✅ Matched Keywords**")
                for kw in ats_analysis.get('matched_keywords', [])[:10]:
                    st.markdown(f"• {kw}")
            
            with analysis_col2:
                st.markdown("**❌ Missing Keywords**")
                for kw in ats_analysis.get('missing_keywords', [])[:10]:
                    st.markdown(f"• {kw}")
        
        # AI Recommendations
        if 'ai_suggestions' in st.session_state:
            with st.expander("💡 AI-Powered Improvement Suggestions"):
                suggestions = st.session_state['ai_suggestions']
                
                suggestion_tabs = st.tabs(["🚀 Quick Wins", "📝 Content", "🔧 Technical", "📈 Long-term"])
                
                with suggestion_tabs[0]:
                    st.markdown("**Quick Fixes (Immediate Impact)**")
                    for suggestion in suggestions.get('quick_wins', []):
                        st.markdown(f"⚡ {suggestion}")
                
                with suggestion_tabs[1]:
                    st.markdown("**Content Improvements**")
                    for suggestion in suggestions.get('content_improvements', []):
                        st.markdown(f"📝 {suggestion}")
                
                with suggestion_tabs[2]:
                    st.markdown("**Technical Enhancements**")
                    for suggestion in suggestions.get('keyword_optimization', []):
                        st.markdown(f"🔧 {suggestion}")
                    for suggestion in suggestions.get('formatting_enhancements', []):
                        st.markdown(f"🎨 {suggestion}")
                
                with suggestion_tabs[3]:
                    st.markdown("**Long-term Development Goals**")
                    for goal in suggestions.get('long_term_goals', []):
                        st.markdown(f"📈 {goal}")
        
        # AI-Generated Cover Letter Section
        st.markdown('<h2 class="sub-header">📝 AI Cover Letter Generator</h2>', unsafe_allow_html=True)
        
        if st.button("✨ Generate AI Cover Letter", use_container_width=True):
            with st.spinner("🤖 AI is writing your personalized cover letter..."):
                cover_letter = optimizer.ai_generate_cover_letter(
                    st.session_state['optimized_resume'],
                    st.session_state.get('job_requirements', {}),
                    st.session_state.get('job_description', ''),
                    st.session_state.get('company_name', '')
                )
                st.session_state['cover_letter'] = cover_letter
        
        if 'cover_letter' in st.session_state:
            st.markdown("### 📄 AI-Generated Cover Letter")
            with st.container():
                st.markdown(
                    f'<div style="background-color: #f8f9fa; padding: 1.5rem; border-radius: 0.5rem; border-left: 4px solid #007bff;">'
                    f'{st.session_state["cover_letter"].replace(chr(10), "<br>")}'
                    f'</div>',
                    unsafe_allow_html=True
                )
            
            # Cover letter download
            st.download_button(
                label="📥 Download Cover Letter",
                data=st.session_state['cover_letter'],
                file_name="ai_cover_letter.txt",
                mime="text/plain",
                use_container_width=True
            )
        
        # Download Section
        st.markdown('<h2 class="sub-header">⬇️ Download Optimized Resume</h2>', unsafe_allow_html=True)
        
        download_col1, download_col2, download_col3 = st.columns(3)
        
        with download_col1:
            if st.button("📄 Generate PDF", type="primary", use_container_width=True):
                with st.spinner("📝 Creating professional PDF..."):
                    try:
                        pdf_buffer = optimizer.generate_pdf_resume(st.session_state['optimized_resume'])
                        
                        st.download_button(
                            label="📥 Download PDF Resume",
                            data=pdf_buffer.getvalue(),
                            file_name="ai_optimized_resume.pdf",
                            mime="application/pdf",
                            use_container_width=True,
                            key="download_pdf"
                        )
                        st.success("✅ PDF generated successfully!")
                    except Exception as e:
                        st.error(f"❌ Error generating PDF: {str(e)}")
        
        with download_col2:
            json_data = json.dumps(st.session_state['optimized_resume'], indent=2)
            st.download_button(
                label="📋 Download JSON Data",
                data=json_data,
                file_name="resume_data.json",
                mime="application/json",
                use_container_width=True,
                key="download_json"
            )
        
        with download_col3:
            # Generate comprehensive report
            if st.button("📊 Generate AI Report", use_container_width=True):
                report_data = generate_ai_analysis_report(
                    st.session_state.get('resume_data', {}),
                    st.session_state.get('optimized_resume', {}),
                    st.session_state.get('ats_analysis', {}),
                    st.session_state.get('ai_suggestions', {}),
                    st.session_state.get('job_description', ''),
                    st.session_state.get('company_name', '')
                )
                
                st.download_button(
                    label="📥 Download AI Analysis Report",
                    data=report_data,
                    file_name="ai_resume_analysis_report.txt",
                    mime="text/plain",
                    use_container_width=True,
                    key="download_report"
                )
        
        # Comparison Dashboard
        st.markdown('<h2 class="sub-header">📊 Before vs After AI Optimization</h2>', unsafe_allow_html=True)
        
        comparison_col1, comparison_col2 = st.columns(2)
        
        with comparison_col1:
            st.markdown("### 📄 Original Resume")
            original_metrics = calculate_resume_metrics(st.session_state.get('resume_data', {}))
            st.metric("Structure Score", f"{original_metrics['structure_score']}%")
            st.metric("Skills Count", original_metrics['skills_count'])
            st.metric("Sections Present", f"{original_metrics['sections']}/8")
        
        with comparison_col2:
            st.markdown("### ✨ AI-Optimized Resume")
            optimized_metrics = calculate_resume_metrics(st.session_state.get('optimized_resume', {}))
            improvement = ats_analysis.get('overall_score', 0) - original_metrics.get('basic_score', 0)
            st.metric("AI ATS Score", f"{ats_analysis.get('overall_score', 0)}%", delta=f"+{improvement:.1f}%")
            st.metric("Enhanced Skills", optimized_metrics['skills_count'])
            st.metric("Complete Sections", f"{optimized_metrics['sections']}/8")
        
        # AI Insights Summary
        st.markdown('<h2 class="sub-header">🧠 AI Insights Summary</h2>', unsafe_allow_html=True)
        
        insights_col1, insights_col2, insights_col3 = st.columns(3)
        
        with insights_col1:
            st.markdown("### 🎯 **Keyword Optimization**")
            matched_count = len(ats_analysis.get('matched_keywords', []))
            total_keywords = matched_count + len(ats_analysis.get('missing_keywords', []))
            match_rate = (matched_count / max(total_keywords, 1)) * 100
            st.metric("Keyword Match Rate", f"{match_rate:.1f}%")
            if match_rate >= 80:
                st.success("🟢 Excellent keyword coverage!")
            elif match_rate >= 60:
                st.warning("🟡 Good keyword coverage")
            else:
                st.error("🔴 Needs more keyword optimization")
        
        with insights_col2:
            st.markdown("### 🛠️ **Skills Analysis**")
            skills_score = ats_analysis.get('skills_score', 0)
            st.metric("Skills Alignment", f"{skills_score}%")
            if skills_score >= 85:
                st.success("🟢 Skills perfectly aligned!")
            elif skills_score >= 70:
                st.warning("🟡 Good skills alignment")
            else:
                st.error("🔴 Skills need enhancement")
        
        with insights_col3:
            st.markdown("### 📈 **Overall Assessment**")
            overall_score = ats_analysis.get('overall_score', 0)
            st.metric("Optimization Success", f"{overall_score}%")
            if overall_score >= 90:
                st.success("🟢 Outstanding optimization!")
            elif overall_score >= 75:
                st.warning("🟡 Good optimization")
            else:
                st.error("🔴 Needs further optimization")
        
        # Action Items
        st.markdown("### 🎯 Next Steps")
        next_steps_col1, next_steps_col2 = st.columns(2)
        
        with next_steps_col1:
            st.markdown("**Immediate Actions:**")
            if overall_score < 85:
                st.markdown("• Review AI suggestions above")
                st.markdown("• Incorporate missing keywords")
                st.markdown("• Enhance experience descriptions")
            else:
                st.markdown("• ✅ Resume is well-optimized!")
                st.markdown("• Consider cover letter customization")
                st.markdown("• Prepare for interviews")
        
        with next_steps_col2:
            st.markdown("**Optimization Tips:**")
            st.markdown("• Use quantified achievements")
            st.markdown("• Match job description language")
            st.markdown("• Keep formatting ATS-friendly")
            st.markdown("• Update for each application")

def calculate_resume_metrics(resume_data):
    """Calculate comprehensive resume metrics"""
    metrics = {
        'structure_score': 0,
        'skills_count': 0,
        'sections': 0,
        'basic_score': 0
    }
    
    if not resume_data:
        return metrics
    
    # Count sections
    sections = ['personal_info', 'summary', 'experience', 'skills', 'education', 'projects', 'certifications', 'achievements']
    metrics['sections'] = sum([1 for section in sections if resume_data.get(section)])
    
    # Count skills
    skills = resume_data.get('skills', {})
    if isinstance(skills, dict):
        all_skills = []
        for skill_list in skills.values():
            if isinstance(skill_list, list):
                all_skills.extend(skill_list)
        metrics['skills_count'] = len(all_skills)
    elif isinstance(skills, list):
        metrics['skills_count'] = len(skills)
    
    # Structure score
    metrics['structure_score'] = min(100, (metrics['sections'] / 8) * 100)
    
    # Basic score
    basic_score = 0
    personal_info = resume_data.get('personal_info', {})
    if personal_info.get('name'): basic_score += 10
    if personal_info.get('email'): basic_score += 10
    if personal_info.get('phone'): basic_score += 10
    if resume_data.get('summary'): basic_score += 15
    if resume_data.get('experience'): basic_score += 25
    if resume_data.get('skills'): basic_score += 15
    if resume_data.get('education'): basic_score += 10
    if resume_data.get('projects'): basic_score += 5
    
    metrics['basic_score'] = basic_score
    
    return metrics

def generate_ai_analysis_report(original_resume, optimized_resume, ats_analysis, ai_suggestions, job_description, company_name):
    """Generate comprehensive AI analysis report"""
    
    report = f"""
═══════════════════════════════════════════════════════════════════
                    AI-POWERED RESUME ANALYSIS REPORT
═══════════════════════════════════════════════════════════════════
Generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}
Company: {company_name or 'Not specified'}

EXECUTIVE SUMMARY
═══════════════════════════════════════════════════════════════════
Overall ATS Score: {ats_analysis.get('overall_score', 0)}%
Optimization Status: {'EXCELLENT' if ats_analysis.get('overall_score', 0) >= 90 else 'GOOD' if ats_analysis.get('overall_score', 0) >= 75 else 'NEEDS IMPROVEMENT'}

KEY PERFORMANCE INDICATORS
═══════════════════════════════════════════════════════════════════
• Keyword Matching: {ats_analysis.get('keyword_score', 0)}%
• Skills Alignment: {ats_analysis.get('skills_score', 0)}%
• Experience Relevance: {ats_analysis.get('experience_score', 0)}%
• Education Match: {ats_analysis.get('education_score', 0)}%
• Structure & Format: {ats_analysis.get('structure_score', 0)}%
• Achievement Quantification: {ats_analysis.get('achievement_score', 0)}%

AI-IDENTIFIED STRENGTHS
═══════════════════════════════════════════════════════════════════
"""
    
    for i, strength in enumerate(ats_analysis.get('strengths', []), 1):
        report += f"{i}. {strength}\n"
    
    report += f"""
AI-IDENTIFIED IMPROVEMENT AREAS
═══════════════════════════════════════════════════════════════════
"""
    
    for i, improvement in enumerate(ats_analysis.get('improvements', []), 1):
        report += f"{i}. {improvement}\n"
    
    report += f"""
KEYWORD ANALYSIS
═══════════════════════════════════════════════════════════════════
Matched Keywords ({len(ats_analysis.get('matched_keywords', []))}):
{', '.join(ats_analysis.get('matched_keywords', [])[:20])}

Missing Keywords ({len(ats_analysis.get('missing_keywords', []))}):
{', '.join(ats_analysis.get('missing_keywords', [])[:20])}

AI OPTIMIZATION RECOMMENDATIONS
═══════════════════════════════════════════════════════════════════

QUICK WINS (Immediate Implementation):
"""
    
    for i, suggestion in enumerate(ai_suggestions.get('quick_wins', []), 1):
        report += f"{i}. {suggestion}\n"
    
    report += f"""
CONTENT IMPROVEMENTS:
"""
    
    for i, suggestion in enumerate(ai_suggestions.get('content_improvements', []), 1):
        report += f"{i}. {suggestion}\n"
    
    report += f"""
TECHNICAL ENHANCEMENTS:
"""
    
    for i, suggestion in enumerate(ai_suggestions.get('keyword_optimization', []), 1):
        report += f"{i}. {suggestion}\n"
    
    report += f"""
LONG-TERM DEVELOPMENT GOALS:
"""
    
    for i, goal in enumerate(ai_suggestions.get('long_term_goals', []), 1):
        report += f"{i}. {goal}\n"
    
    report += f"""
ATS RECOMMENDATIONS
═══════════════════════════════════════════════════════════════════
"""
    
    for i, rec in enumerate(ats_analysis.get('ats_recommendations', []), 1):
        report += f"{i}. {rec}\n"
    
    # Resume structure comparison
    original_metrics = calculate_resume_metrics(original_resume)
    optimized_metrics = calculate_resume_metrics(optimized_resume)
    
    report += f"""
RESUME TRANSFORMATION METRICS
═══════════════════════════════════════════════════════════════════
                        BEFORE    AFTER    IMPROVEMENT
Structure Score:        {original_metrics['structure_score']:.1f}%     {optimized_metrics['structure_score']:.1f}%     +{optimized_metrics['structure_score'] - original_metrics['structure_score']:.1f}%
Skills Count:           {original_metrics['skills_count']}        {optimized_metrics['skills_count']}        +{optimized_metrics['skills_count'] - original_metrics['skills_count']}
Complete Sections:      {original_metrics['sections']}/8       {optimized_metrics['sections']}/8       +{optimized_metrics['sections'] - original_metrics['sections']}

IMPLEMENTATION ROADMAP
═══════════════════════════════════════════════════════════════════
Phase 1 (Immediate - 1-2 days):
• Implement quick wins from AI suggestions
• Incorporate high-priority missing keywords according to job description
• Review and enhance quantified achievements

Phase 2 (Short-term - 1 week):
• Refine experience descriptions with action verbs
• Optimize skills section organization
• Enhance professional summary

Phase 3 (Medium-term - 2-4 weeks):
• Develop additional relevant projects
• Obtain missing certifications if applicable
• Build portfolio of quantified achievements

Phase 4 (Long-term - 1-3 months):
• Pursue skill development in missing areas
• Gain experience in highlighted technologies
• Build thought leadership in target domain

QUALITY ASSURANCE CHECKLIST
═══════════════════════════════════════════════════════════════════
□ All contact information is current and professional
□ Resume is saved in ATS-friendly format (PDF)
□ Keywords from job description are naturally incorporated
□ All achievements are quantified where possible
□ Professional summary aligns with job requirements
□ Skills section prioritizes job-relevant competencies
□ Experience descriptions use strong action verbs
□ Resume length is appropriate (1-2 pages)
□ Formatting is clean and ATS-compatible
□ Content is truthful and verifiable

FINAL RECOMMENDATIONS
═══════════════════════════════════════════════════════════════════
Based on AI analysis, your resume shows {'strong potential' if ats_analysis.get('overall_score', 0) >= 75 else 'significant improvement opportunities'}.
Priority focus areas: {', '.join(ats_analysis.get('improvements', [])[:3])}

Next Steps:
1. Implement immediate AI recommendations
2. Customize for each specific job application
3. Prepare compelling interview narratives
4. Continue professional development in identified areas

═══════════════════════════════════════════════════════════════════
Report generated by AI-Powered ATS Resume Builder
For questions or additional optimization, re-run the AI analysis.
═══════════════════════════════════════════════════════════════════
"""
    
    return report

def display_resume_preview(resume_data):
    """Display formatted resume preview with enhanced styling"""
    if not resume_data:
        st.info("No resume data to display")
        return
    
    st.markdown('<div class="resume-preview">', unsafe_allow_html=True)
    
    # Personal Info Header
    personal_info = resume_data.get('personal_info', {})
    if personal_info.get('name'):
        st.markdown(f"# {personal_info['name']}")
        st.markdown("---")
    
    # Contact Information
    contact_info = []
    contact_fields = {
        'email': '📧',
        'phone': '📞', 
        'location': '📍',
        'linkedin': '💼',
        'website': '🌐',
        'github': '🔗'
    }
    
    for field, icon in contact_fields.items():
        if personal_info.get(field):
            if field in ['linkedin', 'website', 'github']:
                contact_info.append(f"{icon} [{field.title()}]({personal_info[field]})")
            else:
                contact_info.append(f"{icon} {personal_info[field]}")
    
    if contact_info:
        st.markdown(" | ".join(contact_info))
        st.markdown("---")
    
    # Professional Summary
    if resume_data.get('summary'):
        st.markdown("## 📝 Professional Summary")
        st.markdown(resume_data['summary'])
        st.markdown("")
    
    # Core Competencies/Skills
    skills = resume_data.get('skills', {})
    if skills:
        st.markdown("## 🛠️ Core Competencies")
        
        if isinstance(skills, dict):
            for category, skill_list in skills.items():
                if skill_list:
                    category_display = category.replace('_', ' ').title()
                    skills_text = " • ".join(skill_list)
                    st.markdown(f"**{category_display}:** {skills_text}")
        elif isinstance(skills, list):
            skills_text = " • ".join(skills)
            st.markdown(skills_text)
        st.markdown("")
    
    # Professional Experience
    if resume_data.get('experience'):
        st.markdown("## 💼 Professional Experience")
        for i, exp in enumerate(resume_data['experience']):
            if exp.get('title'):
                st.markdown(f"### {exp['title']}")
            
            exp_details = []
            if exp.get('company'):
                exp_details.append(f"**{exp['company']}**")
            if exp.get('location'):
                exp_details.append(exp['location'])
            if exp.get('duration'):
                exp_details.append(exp['duration'])
            
            if exp_details:
                st.markdown(" | ".join(exp_details))
            
            if exp.get('description'):
                if isinstance(exp['description'], list):
                    for desc in exp['description']:
                        if desc.strip():
                            st.markdown(f"• {desc}")
                else:
                    st.markdown(exp['description'])
            
            if i < len(resume_data['experience']) - 1:
                st.markdown("---")
        st.markdown("")
    
    # Education
    if resume_data.get('education'):
        st.markdown("## 🎓 Education")
        for edu in resume_data['education']:
            if edu.get('degree'):
                st.markdown(f"**{edu['degree']}**")
            if edu.get('school'):
                school_info = edu['school']
                if edu.get('location'):
                    school_info += f" | {edu['location']}"
                st.markdown(f"*{school_info}*")
            if edu.get('year'):
                st.markdown(f"📅 {edu['year']}")
            if edu.get('gpa'):
                st.markdown(f"**GPA:** {edu['gpa']}")
            if edu.get('honors'):
                st.markdown(f"**Honors:** {edu['honors']}")
            st.markdown("")
    
    # Key Projects
    if resume_data.get('projects'):
        st.markdown("## 🚀 Key Projects")
        for project in resume_data['projects']:
            if project.get('name'):
                project_title = project['name']
                if project.get('url'):
                    project_title = f"[{project_title}]({project['url']})"
                st.markdown(f"### {project_title}")
            
            if project.get('duration'):
                st.markdown(f"📅 {project['duration']}")
            if project.get('description'):
                st.markdown(project['description'])
            if project.get('technologies'):
                if isinstance(project['technologies'], list):
                    tech_text = ", ".join(project['technologies'])
                else:
                    tech_text = project['technologies']
                st.markdown(f"**Technologies:** {tech_text}")
            st.markdown("")
    
    # Certifications
    if resume_data.get('certifications'):
        st.markdown("## 📜 Certifications")
        for cert in resume_data['certifications']:
            if isinstance(cert, dict):
                cert_text = cert.get('name', '')
                if cert.get('issuer'):
                    cert_text += f" - {cert['issuer']}"
                if cert.get('date'):
                    cert_text += f" ({cert['date']})"
                st.markdown(f"• {cert_text}")
            else:
                st.markdown(f"• {cert}")
        st.markdown("")
    
    # Languages
    if resume_data.get('languages'):
        st.markdown("## 🌐 Languages")
        for lang in resume_data['languages']:
            if isinstance(lang, dict):
                lang_text = f"{lang.get('language', '')} - {lang.get('proficiency', '')}"
            else:
                lang_text = lang
            st.markdown(f"• {lang_text}")
        st.markdown("")
    
    # Achievements
    if resume_data.get('achievements'):
        st.markdown("## 🏆 Key Achievements")
        for achievement in resume_data['achievements']:
            st.markdown(f"• {achievement}")
        st.markdown("")
    
    st.markdown('</div>', unsafe_allow_html=True)

def add_footer():
    """Add a beautiful footer to the app"""
    st.markdown("""
    <div class="footer">
        Created with <span class="footer-heart">❤️</span> for <span class="footer-name">Nagashree</span>
    </div>
    """, unsafe_allow_html=True)

add_footer() 

if __name__ == "__main__":
    main()